import io
import os
import re
from typing import List

import streamlit as st
from google import genai
from pydantic import BaseModel, Field
from pypdf import PdfReader
from docx import Document

MODEL_NAME = "gemini-3.7-flash"
MAX_RESUME_CHARS = 50000


class Improvement(BaseModel):
    area: str
    priority: str = Field(description="High, Medium, or Low")
    issue: str
    recommendation: str
    example: str


class ATSResult(BaseModel):
    ats_score: int = Field(ge=0, le=100)
    verdict: str
    summary: str
    strengths: List[str]
    improvements: List[Improvement]
    matched_keywords: List[str]
    missing_keywords: List[str]
    formatting_issues: List[str]
    section_feedback: List[str]


def extract_text(uploaded_file) -> str:
    """Extract text from PDF, DOCX, or TXT uploads."""
    name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()

    if name.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()

    if name.endswith(".docx"):
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts).strip()

    if name.endswith(".txt"):
        return data.decode("utf-8", errors="replace").strip()

    raise ValueError("Unsupported file type. Upload PDF, DOCX, or TXT.")


def clean_resume_text(text: str) -> str:
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text[:MAX_RESUME_CHARS].strip()


def analyze_resume(resume_text: str, job_description: str, api_key: str) -> ATSResult:
    client = genai.Client(api_key=api_key)
    jd_context = (
        job_description.strip()
        or "No job description was supplied. Evaluate general ATS readiness and clarity "
           "rather than job-specific matching."
    )

    prompt = f"""
You are an ATS resume evaluator and senior technical recruiter.
Evaluate the resume below for ATS compatibility and recruiter usefulness.

Important:
- This is a simulated ATS-readiness score, not a score from a real ATS vendor.
- If a job description is supplied, prioritize relevance to it and identify missing keywords.
- Never invent experience, education, certifications, or skills that are not present.
- Penalize tables/columns only when they are explicitly indicated by extracted text or obvious
  structural clues; do not assume formatting that text extraction cannot reveal.
- Focus on standard section headings, measurable achievements, keyword alignment, clarity,
  chronology, and ATS-safe wording.
- Give practical edits, including before/after examples when possible.
- Score 0-100 using this rough rubric: keyword/relevance 35, experience/achievement quality 25,
  structure/section completeness 20, ATS-safe clarity 10, grammar/consistency 10.

JOB DESCRIPTION:
{jd_context}

RESUME:
{resume_text}
"""

    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=prompt,
        config={
            "temperature": 0.2,
            "response_mime_type": "application/json",
            "response_schema": ATSResult.model_json_schema(),
        },
    )

    if not response.text:
        raise RuntimeError("Gemini returned an empty response.")

    try:
        return ATSResult.model_validate_json(response.text)
    except Exception as exc:
        raise RuntimeError(
            f"Could not parse Gemini's structured response: {exc}"
        ) from exc


def get_api_key() -> str:
    try:
        secret_key = st.secrets.get("GEMINI_API_KEY", "")
    except Exception:
        secret_key = ""
    return secret_key or os.getenv("GEMINI_API_KEY", "")


def main() -> None:
    st.set_page_config(
        page_title="Resume ATS Analyzer",
        page_icon="📄",
        layout="wide",
    )
    st.title("📄 Resume ATS Analyzer")
    st.caption(
        "Upload a resume to get a simulated ATS-readiness score and actionable improvements."
    )

    with st.sidebar:
        st.header("Settings")
        st.markdown(f"**Model:** `{MODEL_NAME}`")
        st.info(
            "For best results, add the target job description. The score is an AI-based estimate, "
            "not an official ATS score."
        )

    uploaded_file = st.file_uploader(
        "Upload your resume",
        type=["pdf", "docx", "txt"],
    )
    job_description = st.text_area(
        "Target job description (optional)",
        height=220,
        placeholder=(
            "Paste the job description here to get job-specific keyword matching "
            "and recommendations."
        ),
    )

    if uploaded_file is None:
        st.markdown("### What you'll get")
        st.markdown(
            "- **0–100 ATS-readiness score**\n"
            "- **Keyword matches and gaps**\n"
            "- **Prioritized improvements**\n"
            "- **Formatting and section feedback**\n"
            "- **Specific rewrite examples**"
        )
        return

    if st.button("Analyze Resume", type="primary", use_container_width=True):
        api_key = get_api_key()
        if not api_key:
            st.error(
                "Gemini API key not found. Add GEMINI_API_KEY to Streamlit secrets "
                "or your environment."
            )
            st.stop()

        try:
            with st.spinner("Extracting resume and analyzing with Gemini…"):
                resume_text = clean_resume_text(extract_text(uploaded_file))
                if len(resume_text) < 100:
                    st.error(
                        "Very little text was extracted. Try a text-based PDF/DOCX or a TXT file."
                    )
                    st.stop()
                result = analyze_resume(resume_text, job_description, api_key)
        except Exception as exc:
            st.error(f"Analysis failed: {exc}")
            st.stop()

        st.success("Analysis complete.")

        col1, col2, col3 = st.columns(3)
        col1.metric("ATS readiness", f"{result.ats_score}/100")
        col2.metric("Matched keywords", len(result.matched_keywords))
        col3.metric("Missing keywords", len(result.missing_keywords))

        st.progress(result.ats_score / 100)

        st.subheader("Verdict")
        st.write(result.verdict)

        st.subheader("Summary")
        st.write(result.summary)

        left, right = st.columns(2)

        with left:
            st.subheader("Strengths")
            for item in result.strengths:
                st.markdown(f"- {item}")

            st.subheader("Matched keywords")
            st.write(
                ", ".join(result.matched_keywords)
                if result.matched_keywords
                else "None identified"
            )

        with right:
            st.subheader("Missing keywords")
            st.write(
                ", ".join(result.missing_keywords)
                if result.missing_keywords
                else "None identified"
            )

            st.subheader("Formatting issues")
            if result.formatting_issues:
                for item in result.formatting_issues:
                    st.markdown(f"- {item}")
            else:
                st.write("No major issues identified from the extracted text.")

        st.subheader("Prioritized improvements")
        for item in result.improvements:
            with st.expander(f"{item.priority} priority — {item.area}"):
                st.write(f"**Issue:** {item.issue}")
                st.write(f"**Recommendation:** {item.recommendation}")
                if item.example:
                    st.write(f"**Example:** {item.example}")

        st.subheader("Section feedback")
        for item in result.section_feedback:
            st.markdown(f"- {item}")

        with st.expander("View extracted resume text"):
            st.text(resume_text)


if __name__ == "__main__":
    main()
